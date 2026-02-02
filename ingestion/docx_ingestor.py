# ingestion/docx_ingestor.py
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocxIngestor:
    def __init__(self, docx_dir: str = "data/docx"):
        self.docx_dir = Path(docx_dir)

    def ingest(self) -> List[Document]:
        documents: List[Document] = []

        for docx_path in self.docx_dir.glob("*.docx"):
            doc = DocxDocument(docx_path)
            blocks: List[str] = []

            # Paragraphs
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    blocks.append(text)

            # Tables 
            for table in doc.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if cells:
                        blocks.append(" | ".join(cells))

            full_text = "\n".join(blocks).strip()

            # Prevent garbage embeddings
            if len(full_text) < 200:
                continue

            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "document_type": "docx",
                        "source": docx_path.name,
                        "section": "full_document",
                    },
                )
            )

        return documents
